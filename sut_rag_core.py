# sut_rag_core.py
# Description: SUT RAG Engine with Deep Chunking, FTS, Hybrid Search, and Golden Loop.
# Status: Production Ready (Final)

import os
import re as regex
import json
import uuid
import sqlite3
import numpy as np
import faiss
import pypandoc
from typing import List, Dict, Generator
from sentence_transformers import CrossEncoder 
# LangChain & AI Libraries
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import MarkdownHeaderTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage

# Optional Imports (with safety checks)
try:
    from docx import Document
except ImportError:
    print("[WARN] 'python-docx' library not found.")
try:
    from thefuzz import fuzz
except ImportError:
    pass

# --- Configuration ---
DOCX_FILE_PATH = "08.03.2025-Değişiklik Tebliği İşlenmiş Güncel 2013 SUT.docx"
MARKDOWN_FILE_PATH = "sut_converted_temp.md"
DB_PATH = "sut_knowledge_base.db"
FAISS_INDEX_PATH = "sut_faiss.index"
FAISS_MAPPING_PATH = "sut_faiss.index.mapping"
EMBEDDING_MODEL = "paraphrase-multilingual-MiniLM-L12-v2"


class SUT_RAG_Engine:
    def __init__(self, llm_provider: str = "google", model_name: str = "gemini-2.5-flash"):
        self.embeddings_model = self._initialize_embeddings()
        print("[INIT] Loading Reranker Model...")
        self.reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2', device='cpu')
        self.conn = None
        self.cursor = None
        self.faiss_index = None
        self.id_mapping = None
        self.llm = None
        self.provider = llm_provider

        print(f"[INIT] SUT Engine Initialized. Provider: '{llm_provider}', Model: '{model_name}'")

        if llm_provider == "google":
            self._init_google_llm(model_name)
        elif llm_provider == "openrouter":
            self._init_openrouter_llm(model_name)
        elif llm_provider == "lmstudio":
            self._init_lmstudio_llm(model_name)
        else:
            print(f"[ERROR] Unsupported provider: {llm_provider}")

    # --- LLM Initialization ---
    def _init_google_llm(self, model_name: str):
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            print("[ERROR] GEMINI_API_KEY missing in .env")
            return
        # Lower temperature for factual consistency
        self.llm = ChatGoogleGenerativeAI(model=model_name, api_key=api_key, temperature=0.1)

    def _init_openrouter_llm(self, model_name: str):
        api_key = os.getenv("OPENROUTER_API_KEY")
        if not api_key: return
        self.llm = ChatOpenAI(
            model=model_name, openai_api_key=api_key, openai_api_base="https://openrouter.ai/api/v1",
            temperature=0.1
        )

    def _init_lmstudio_llm(self, model_name: str):
        self.llm = ChatOpenAI(
            model=model_name, openai_api_key="not-needed", openai_api_base="http://localhost:1234/v1",
            temperature=0.1
        )

    def _initialize_embeddings(self):
        return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL, model_kwargs={'device': 'cpu'})

    def __del__(self):
        if self.conn: self.conn.close()

    # =========================================================================
    # MODIFICATION 1 & 2 & 3: Robust Prep, Deep Chunking, FTS Database
    # =========================================================================

    def populate_database(self):
        """Orchestrates the full database creation pipeline."""
        print("[PREP] Starting database population...")
        
        # 1. Robust File Cleaning
        cleaned_path = self._remove_strikethrough_and_save_temp(DOCX_FILE_PATH)
        if not cleaned_path:
            print("[ERROR] Failed to clean DOCX. Aborting.")
            return

        # 2. Convert to Markdown
        print("[PREP] Converting to Markdown (Pandoc)...")
        try:
            pypandoc.convert_file(cleaned_path, 'md', outputfile=MARKDOWN_FILE_PATH)
        except Exception as e:
            print(f"[ERROR] Pandoc conversion failed: {e}")
            return

        # 3. Deep Semantic Chunking (H1-H4)
        print("[PREP] Splitting text into semantic chunks...")
        chunks = self._get_markdown_chunks(MARKDOWN_FILE_PATH)
        if not chunks: return

        # 4. Setup Database (with FTS)
        self._setup_database()

        print("[DB] Inserting data into SQLite (Standard + FTS) and FAISS...")
        texts_to_embed, string_ids = [], []

        for chunk in chunks:
            chunk_id = str(uuid.uuid4())
            metadata_json = json.dumps(chunk.metadata, ensure_ascii=False)
            page_content = chunk.page_content
            
            # Extract headers for FTS (Full Text Search)
            header_text = " ".join([v for k, v in chunk.metadata.items() if k.startswith("Header")])

            # Insert into Standard Table
            self.cursor.execute(
                "INSERT INTO chunks (chunk_id, text_content, metadata_json) VALUES (?, ?, ?)",
                (chunk_id, page_content, metadata_json)
            )

            # Insert into FTS Virtual Table (Mod 3)
            self.cursor.execute(
                "INSERT INTO title_search (chunk_id, header_text) VALUES (?, ?)",
                (chunk_id, header_text)
            )

            # Prepare for FAISS
            full_text_for_embed = f"{header_text}\n\n{page_content}"
            texts_to_embed.append(full_text_for_embed)
            string_ids.append(chunk_id)

        self.conn.commit()

        # 5. Create FAISS Index
        self._create_and_save_faiss_index(texts_to_embed, string_ids)
        
        # Cleanup
        if os.path.exists(cleaned_path): os.remove(cleaned_path)
        if os.path.exists(MARKDOWN_FILE_PATH): os.remove(MARKDOWN_FILE_PATH)
        print("[SUCCESS] Database population complete.")

    def _remove_strikethrough_and_save_temp(self, input_path):
        """(Mod 1) Robustly removes strikethrough text."""
        if not os.path.exists(input_path):
            print(f"[ERROR] File not found: {input_path}")
            return None
            
        try:
            doc = Document(input_path)
            temp_output = "temp_cleaned_sut.docx"
            
            # Helper to process paragraphs
            def clean_p(paragraph):
                # Iterate backwards to safely remove runs without messing up indices
                for i in range(len(paragraph.runs)-1, -1, -1):
                    run = paragraph.runs[i]
                    # Check both direct formatting and style formatting
                    if run.font.strike:
                        paragraph._p.remove(run._r)
            
            for p in doc.paragraphs: clean_p(p)
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: clean_p(p)
            
            doc.save(temp_output)
            return temp_output
        except Exception as e:
            print(f"[ERROR] DOCX processing failed: {e}")
            return None

    def _get_markdown_chunks(self, md_path):
        """(Mod 2) Deep Chunking handling H1 to H4 hierarchy."""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Cleaning artifacts
            text = regex.sub(r'~~.*?~~', '', text)
            text = regex.sub(r'►', '', text)

            # Header Normalization: Convert "1.2.3.4 Title" to "#### 1.2.3.4 Title"
            def h_repl(m):
                depth = m.group(1).count('.') + 1
                hashes = '#' * min(6, depth) # Markdown supports up to H6
                return f"{hashes} {m.group(0)}"
            
            # Apply regex for deep headers
            text = regex.sub(r"^\*\*((\d+\.)+\d+[\.\d\w-]*)\s*-*\s*([^ \n\*]+.*?)\*\*", h_repl, text, flags=regex.MULTILINE)

            # Deep Splitting Configuration
            headers_to_split_on = [
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
                ("####", "Header 4"),
            ]
            
            splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
            return splitter.split_text(text)
        except Exception as e:
            print(f"[ERROR] Chunking failed: {e}")
            return []

    def _setup_database(self):
        """(Mod 3) Database Setup with FTS5."""
        if os.path.exists(DB_PATH): os.remove(DB_PATH)
        if os.path.exists(FAISS_INDEX_PATH): os.remove(FAISS_INDEX_PATH)
        if os.path.exists(FAISS_MAPPING_PATH): os.remove(FAISS_MAPPING_PATH)

        self.conn = sqlite3.connect(DB_PATH, check_same_thread=False)
        self.cursor = self.conn.cursor()

        # Standard Table
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS chunks (
                chunk_id TEXT PRIMARY KEY,
                text_content TEXT NOT NULL,
                metadata_json TEXT
            )
        """)
        
        # FTS Virtual Table (Full Text Search)
        # Note: 'fts5' is built-in to most Python SQLite distributions
        try:
            self.cursor.execute("CREATE VIRTUAL TABLE IF NOT EXISTS title_search USING fts5(chunk_id, header_text)")
        except Exception:
            print("[WARN] FTS5 not supported in this SQLite version. Title search may be slower.")

        self.conn.commit()

    def _create_and_save_faiss_index(self, texts, ids):
        vectors = self.embeddings_model.embed_documents(texts)
        index = faiss.IndexFlatL2(len(vectors[0]))
        index.add(np.array(vectors).astype('float32'))
        faiss.write_index(index, FAISS_INDEX_PATH)
        with open(FAISS_MAPPING_PATH, "w", encoding="utf-8") as f:
            json.dump(ids, f)
        self.faiss_index = index
        self.id_mapping = ids

    # =========================================================================
    # MODIFICATION 4: Hybrid Retrieval & Query Expansion
    # =========================================================================

    def load_database(self) -> bool:
        if not os.path.exists(FAISS_INDEX_PATH): return False
        self.conn = sqlite3.connect(DB_PATH, check_same_thread=False)
        self.cursor = self.conn.cursor()
        self.faiss_index = faiss.read_index(FAISS_INDEX_PATH)
        with open(FAISS_MAPPING_PATH, "r", encoding="utf-8") as f:
            self.id_mapping = json.load(f)
        return True
    def _expand_query(self, user_query: str) -> str:
        """Uses LLM to add synonyms (Stricter Prompt)."""
        if not self.llm: return user_query
        
        # Strict prompt to prevent SQL generation
        prompt_text = (
            f"Kullanıcı sorusu: '{user_query}'\n"
            "Görevin: Bu soruyu SUT (Sağlık Uygulama Tebliği) içinde aramak için anahtar kelimelerle genişlet.\n"
            "Çıktı Formatı: Sadece aralarında boşluk olan kelimeler yaz. SQL veya kod yazma.\n"
            "Örnek: 'kanser ilacı' -> 'kanser onkoloji kemoterapi antineoplastik ilaç ödeme koşulları'\n"
            "Genişletilmiş Kelimeler:"
        )
        
        try:
            messages = [HumanMessage(content=prompt_text)]
            response = self.llm.invoke(messages)
            content = response.content if hasattr(response, 'content') else str(response)
            # Cleanup if model still tries to be chatty
            clean_content = content.replace("SELECT", "").replace("FROM", "").replace("\n", " ")
            return f"{user_query} {clean_content}".strip()
        except Exception:
            return user_query



    def _retrieve_chunks(self, query: str, k: int) -> List[Dict]:
        """(Mod 4) Hybrid Search + Reranking."""
        if not self.faiss_index: return []

        # 1. Vector Search
        # Fetch 3x more candidates than needed for Reranking to filter
        initial_k = k * 3 
        q_vec = np.array([self.embeddings_model.embed_query(query)]).astype('float32')
        _, indices = self.faiss_index.search(q_vec, initial_k)
        
        vector_ids = []
        for idx in indices[0]:
            if idx != -1: vector_ids.append(self.id_mapping[idx])

        # 2. Keyword/Code Search
        keyword_ids = []
        if any(char.isdigit() for char in query):
            clean_term = query.replace("'", "")
            try:
                # FTS Search
                self.cursor.execute("SELECT chunk_id FROM title_search WHERE header_text MATCH ? LIMIT ?", (f'"{clean_term}"', initial_k))
                keyword_ids = [row[0] for row in self.cursor.fetchall()]
            except:
                # LIKE Fallback
                self.cursor.execute("SELECT chunk_id FROM chunks WHERE metadata_json LIKE ? LIMIT ?", (f"%{clean_term}%", initial_k))
                keyword_ids = [row[0] for row in self.cursor.fetchall()]

        # 3. Merge Unique IDs
        combined_ids = list(dict.fromkeys(keyword_ids + vector_ids))
        
        # 4. Fetch Content for Reranking
        candidates = []
        for cid in combined_ids:
            self.cursor.execute("SELECT chunk_id, text_content, metadata_json FROM chunks WHERE chunk_id=?", (cid,))
            row = self.cursor.fetchone()
            if row:
                candidates.append({
                    "id": row[0], 
                    "text": row[1], 
                    "metadata": json.loads(row[2])
                })

        # --- 5. RERANKING STEP ---
        if candidates:
            try:
                # Prepare pairs for CrossEncoder: [[query, text1], [query, text2]]
                pairs = [[query, doc['text']] for doc in candidates]
                
                # Predict scores
                scores = self.reranker.predict(pairs)
                
                # Attach scores to docs
                for doc, score in zip(candidates, scores):
                    doc['score'] = score
                
                # Sort by Score (High to Low)
                candidates.sort(key=lambda x: x['score'], reverse=True)
                
            except Exception as e:
                print(f"[WARN] Reranking failed: {e}. Returning default order.")
        
        # Return only the requested top k
        return candidates[:k]

        # =========================================================================
        # MODIFICATION 5: The "Golden" Agent Loop (Direct Context Chain)
    def query_agentic_rag_stream(self, user_query: str, k: int = 5) -> Generator[Dict, None, None]:
            if self.llm is None:
                yield {"error": "LLM not initialized."}
                return

            yield {"status": "Sorgu genişletiliyor..."}
            expanded_query = self._expand_query(user_query)

            yield {"status": "SUT veritabanı taranıyor..."}
            chunks = self._retrieve_chunks(expanded_query, k)
            
            if not chunks:
                yield {"final_answer": "SUT içinde ilgili bilgi bulunamadı."}
                return

            yield {"status": "Bağlam limiti kontrol ediliyor..."}
            
            # Context Budgeting
            MAX_CONTEXT_CHARS = 12000 
            context_str = ""
            current_char_count = 0
            
            for i, c in enumerate(chunks):
                headers = [v for k, v in c['metadata'].items() if k.startswith("Header")]
                breadcrumb = " > ".join(headers) if headers else "Bölüm"
                chunk_text = c['text']
                if len(chunk_text) > 3500: chunk_text = chunk_text[:3500] + "... [Kısaltıldı]"
                formatted = f"\n--- KAYNAK {i+1} ---\nBAŞLIK: {breadcrumb}\nİÇERİK:\n{chunk_text}\n"
                if current_char_count + len(formatted) > MAX_CONTEXT_CHARS: break
                context_str += formatted
                current_char_count += len(formatted)
                yield {"source": {"title": breadcrumb, "content": chunk_text}}

            system_prompt = f"""
    ROLE:
    You are an expert AI assistant for the Turkish SUT (Healthcare).
    Answer using ONLY the context provided.

    CONTEXT:
    {context_str}

    RULES:
    1. Cite specific article numbers.
    2. If info is missing, say "Bilgi yok".
    3. Respond in Turkish.
    """

            try:
                messages = [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_query)
                ]
                
                # Variables to handle both Delta and Block streaming
                accumulated_response = ""
                accumulated_reasoning = ""
                
                for chunk in self.llm.stream(messages):
                    
                    # 1. Extract Content
                    content = chunk.content if hasattr(chunk, 'content') else ""
                    
                    # 2. Extract Hidden Reasoning (Specific to LM Studio/DeepSeek)
                    reasoning = ""
                    kwargs = getattr(chunk, 'additional_kwargs', {})
                    if kwargs:
                        reasoning = kwargs.get('reasoning', '') or kwargs.get('reasoning_content', '')

                    # Handle Reasoning
                    if reasoning:
                        accumulated_reasoning += reasoning
                        yield {"analysis_content": accumulated_reasoning}

                    # Handle Content
                    if content:
                        accumulated_response += content
                        
                        # Logic: If it contains Harmony tags, use the parser.
                        # If NOT, yield the FULL accumulated text so the UI can overwrite safely.
                        if "<|start|>" in accumulated_response and "channel=analysis" in accumulated_response:
                            for segment in self._parse_harmony_buffer(accumulated_response):
                                yield segment
                        else:
                            # CRITICAL FIX: Sending full accumulated text instead of delta
                            # This fixes the flickering/empty issue in Local Mode
                            yield {"final_answer": accumulated_response}

                yield {"done": True}

            except Exception as e:
                if "context" in str(e).lower():
                    yield {"error": "Model hafızası doldu."}
                else:
                    yield {"error": f"LLM Generation Error: {str(e)}"}

    def _parse_harmony_buffer(self, text: str):
        results = []
        an = regex.search(r"channel=analysis<\|message\|>(.*?)<\|end\|>", text, regex.DOTALL)
        if an: results.append({"analysis_content": an.group(1).strip()})
        
        fi = regex.search(r"channel=final<\|message\|>(.*?)<\|end\|>", text, regex.DOTALL)
        if fi: results.append({"final_answer": fi.group(1).strip()})
        return results