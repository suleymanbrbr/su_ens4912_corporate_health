# rag_storage.py
# Description: Handles SUT Data Extraction, Chunking, and Postgres Storage.

import logging
import os
import re as regex
import uuid
import pypandoc
from typing import List, Dict  # noqa: F401  (kept for downstream typing)
import psycopg2
from psycopg2.extras import Json

logger = logging.getLogger(__name__)

try:
    from docx import Document
except ImportError:
    logger.warning("'python-docx' library not found.")

# --- Configuration ---
DOCX_FILE_PATH = "data/08.03.2025-Değişiklik Tebliği İşlenmiş Güncel 2013 SUT.docx"
MARKDOWN_FILE_PATH = "data/sut_converted_temp.md"

class SUT_Storage_Manager:
    def __init__(self, embeddings_model):
        self.embeddings_model = embeddings_model
        self.conn = None
        self.cursor = None

    def populate_database(self):
        """Orchestrates the full database creation pipeline.

        Always closes the DB connection it opens (even on early-exit / errors)
        so background-task failures don't leak a connection from Neon's pool.
        Raises RuntimeError with a specific step name on failure, so callers
        can surface the exact failing step (HF Space logs are opaque).
        """
        import os as _os
        logger.info("Starting database population...")
        logger.info(f"  cwd      = {_os.getcwd()}")
        logger.info(f"  DOCX     = {DOCX_FILE_PATH}  (exists={_os.path.exists(DOCX_FILE_PATH)})")

        cleaned_path = self._remove_strikethrough_and_save_temp(DOCX_FILE_PATH)
        if not cleaned_path:
            raise RuntimeError(
                f"DOCX cleaning failed (python-docx step). "
                f"Path={DOCX_FILE_PATH}, cwd={_os.getcwd()}, "
                f"exists={_os.path.exists(DOCX_FILE_PATH)}"
            )

        logger.info("Converting to Markdown (Pandoc)...")
        try:
            pypandoc.convert_file(cleaned_path, 'md', outputfile=MARKDOWN_FILE_PATH)
        except Exception as e:
            raise RuntimeError(f"Pandoc conversion failed: {type(e).__name__}: {e}") from e

        logger.info("Splitting text into semantic chunks...")
        chunks = self._get_markdown_chunks(MARKDOWN_FILE_PATH)
        if not chunks:
            raise RuntimeError(
                f"Chunking returned 0 chunks. "
                f"Markdown size={_os.path.getsize(MARKDOWN_FILE_PATH) if _os.path.exists(MARKDOWN_FILE_PATH) else 'NO_FILE'}"
            )

        try:
            self._setup_database()

            logger.info(f"Inserting data into PostgreSQL (pgvector). Total chunks: {len(chunks)}")

            for i, chunk in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                metadata_json = chunk.metadata
                page_content = chunk.page_content
                header_text = " ".join([v for k, v in chunk.metadata.items() if k.startswith("Header")])

                full_text_for_embed = f"{header_text}\n\n{page_content}"
                vector = self.embeddings_model.embed_query(full_text_for_embed)

                self.cursor.execute(
                    "INSERT INTO chunks (chunk_id, text_content, metadata_json, header_text, embedding) VALUES (%s, %s, %s, %s, %s)",
                    (chunk_id, page_content, Json(metadata_json), header_text, vector)
                )

                if (i + 1) % 10 == 0:
                    logger.info(f"Progress: {i + 1}/{len(chunks)} chunks processed.")

            self.conn.commit()
        except Exception as e:
            logger.exception(f"Database population failed: {e}")
            if self.conn:
                try:
                    self.conn.rollback()
                except Exception:  # noqa: BLE001
                    pass
            return False
        finally:
            # Always release the DB connection — even on success the caller
            # never closes it for us.
            if self.cursor is not None:
                try:
                    self.cursor.close()
                except Exception:  # noqa: BLE001
                    pass
            if self.conn is not None:
                try:
                    self.conn.close()
                except Exception:  # noqa: BLE001
                    pass

        if os.path.exists(cleaned_path):
            os.remove(cleaned_path)
        if os.path.exists(MARKDOWN_FILE_PATH):
            os.remove(MARKDOWN_FILE_PATH)
        logger.info("Database population complete.")
        return True

    def _remove_strikethrough_and_save_temp(self, input_path):
        if not os.path.exists(input_path): return None
        try:
            doc = Document(input_path)
            temp_output = "temp_cleaned_sut.docx"
            def clean_p(paragraph):
                for i in range(len(paragraph.runs)-1, -1, -1):
                    run = paragraph.runs[i]
                    if run.font.strike:
                        paragraph._p.remove(run._r)
            for p in doc.paragraphs: clean_p(p)
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: clean_p(p)
            doc.save(temp_output)
            return temp_output
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return None

    def _get_markdown_chunks(self, md_path):
        from langchain_text_splitters import MarkdownHeaderTextSplitter
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                text = f.read()
            text = regex.sub(r'~~.*?~~', '', text)
            text = regex.sub(r'►', '', text)
            def h_repl(m):
                depth = m.group(1).count('.') + 1
                hashes = '#' * min(6, depth)
                return f"{hashes} {m.group(0)}"
            text = regex.sub(r"^\*\*((\d+\.)+\d+[\.\d\w-]*)\s*-*\s*([^ \n\*]+.*?)\*\*", h_repl, text, flags=regex.MULTILINE)
            
            # --- FIX FAKE HEADERS ---
            # Demote any markdown header (#) that does NOT start with a number (like 1., 2.3) or EK-
            def fix_fake_headers(m):
                hashes, title = m.groups()
                title_stripped = title.strip()
                # Remove markdown bold/italic before checking
                clean_title = regex.sub(r'^[\*\s]+', '', title_stripped)
                # If it doesn't start with a digit or 'EK-', demote it to bold text
                if not regex.match(r'^(\d|EK-)', clean_title):
                    return f"**{title_stripped}**"
                return m.group(0)
            
            text = regex.sub(r'^(#{1,6})\s+(.+)$', fix_fake_headers, text, flags=regex.MULTILINE)
            headers_to_split_on = [("#", "Header 1"), ("##", "Header 2"), ("###", "Header 3"), ("####", "Header 4")]
            splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
            md_splits = splitter.split_text(text)
            
            # Since we removed fake headers, some sections might be massive (16k+ chars). 
            # We must chunk them further to fit into the embedding model's 512 token limit.
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            char_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1200, 
                chunk_overlap=200,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            return char_splitter.split_documents(md_splits)
        except Exception as e:
            logger.error(f"Chunking failed: {e}")
            return []

    def _setup_database(self):
        self.conn = psycopg2.connect(os.getenv("DATABASE_URL"))
        self.cursor = self.conn.cursor()
        
        # Purge Knowledge Base Chunks ONLY
        self.cursor.execute("DROP TABLE IF EXISTS chunks CASCADE")
        
        # Enable pgvector if not enabled
        self.cursor.execute("CREATE EXTENSION IF NOT EXISTS vector")

        # RAG Chunks Table
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS chunks (
                chunk_id TEXT PRIMARY KEY,
                text_content TEXT NOT NULL,
                metadata_json JSONB,
                header_text TEXT,
                embedding vector(384)
            )
        """)
        
        # Create FTS Index on header_text and text_content
        self.cursor.execute("CREATE INDEX IF NOT EXISTS chunks_fts_idx ON chunks USING GIN (to_tsvector('turkish', header_text || ' ' || text_content));")

        self.conn.commit()
